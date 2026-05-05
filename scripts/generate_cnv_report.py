import sys
import os
import re
import requests
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import warnings
warnings.simplefilter(action='ignore', category=FutureWarning)

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):

    # Create the relationship ID
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed attributes
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create the w:r element (run)
    new_run = OxmlElement('w:r')

    # Create w:rPr element (run properties)
    rPr = OxmlElement('w:rPr')

    # Color
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), color)
    rPr.append(color_elem)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    else:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    # Set font size (optional)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')  # 10pt
    rPr.append(sz)

    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text

    # Append all together
    new_run.append(rPr)
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return new_run

def extract_name_and_format(s):
    s = s.split('_')[0]

    match = re.search(r'\d', s)
    idx = match.start()
    id_part = s[idx::]

    # Split camel case (e.g., AhmedFathElbab → Ahmed Fath Elbab)
    name_part = s[:idx]
    name_words = re.findall(r'[A-Z][a-z]*', name_part)
    formatted_name = ' '.join(name_words)

    return formatted_name, id_part

def center_table_content(table, horizontal=True, vertical=True):   
    for row in table.rows:
        for cell in row.cells:
            if horizontal:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if vertical:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Set Input and Output files names
input_tsv = sys.argv[1]

file_name = os.path.basename(input_tsv).split(".")[0]
patient_name, patient_id = extract_name_and_format(file_name)

try:
    output_doc = sys.argv[2]
except:
    output_doc = os.path.dirname(input_tsv) + f"./{patient_name} {patient_id} CNV Report.docx"

# Load TSV file into DataFrame
df = pd.read_csv(input_tsv, sep="\t", keep_default_na=False)

if "Report" not in df.columns: # For older versions before adding "Report" column (deprecated)
    df["Report"] = "Primary"

empty_row = pd.DataFrame([[None] * len(df.columns)], columns=df.columns)
primary_variants = df[df["Report"] == "Primary"]
secondary_variants = df[df["Report"] == "Secondary"]
attached_variants = df[df["Report"] == "Attached"]

# Create a new Word Document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


### Header & Personal Information Section ###

doc.add_heading("GENETIC DIAGNOSTIC REPORT", 0)

doc.add_heading('Patient Information', level=1)
table = doc.add_table(rows=2, cols=4)
table.style = 'Light Grid'

# Headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Patient Name'
hdr_cells[1].text = 'Sample ID'
hdr_cells[2].text = 'Sex'
hdr_cells[3].text = 'Service'
# Rows
row_cells = table.rows[1].cells
row_cells[0].text = patient_name
row_cells[1].text = patient_id
row_cells[2].text = '*'
row_cells[3].text = 'WES + CNV'

table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.allow_autofit = True
center_table_content(table)

doc.add_heading("Copy Number Variants", level=1)

### Primary Findings Section ###
if len(primary_variants) > 0:
    doc.add_heading('Primary Findings', level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Type'
    hdr_cells[1].text = 'Zygosity'
    hdr_cells[2].text = 'Variant/Size'
    hdr_cells[3].text = 'Gene'
    hdr_cells[4].text = 'Class'
    hdr_cells[5].text = 'Disease'

    for _, row in pd.concat([primary_variants, empty_row], ignore_index=True).iterrows():
        r = table.add_row().cells
        chr = str(row.get("Chromosome", ""))
        start = str(row.get("Start", ""))
        end = str(row.get("End", ""))
        type = str(row.get("Type", ""))
        classification = str(row.get("Classification", ""))
        genes = str(row.get("Genes", ""))
        omim_genes = str(row.get("OMIM_gene", ""))
        zygosity = str(row.get("Zygosity", ""))
        omim = str(row.get("OMIM", ""))
        size = str(row.get("Length", ""))
        cytoband = str(row.get("Cytoband", ""))

        genes_list = genes.split(", ")
        omim_genes_list = omim_genes.split(", ")
        
        r[0].text = type
        r[1].text = zygosity
        r[2].text = f"chr{chr}:{start}-{end}" + f'\n"{cytoband}"' + f"\n({size} bp)"
        r[3].text = genes
        r[4].text = classification
        r[5].text = "- " + str(omim).replace(" | ", "\n- ") + "\n"
        for i in range(len(genes_list)):
            add_hyperlink(r[5].paragraphs[-1], f"https://www.omim.org/entry/{omim_genes_list[i]}", f"{genes_list[i]} (OMIM: {omim_genes_list[i]}) ,")
            
    last_row = table.rows[len(table.rows) - 1]
    table._tbl.remove(last_row._tr)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    center_table_content(table)

### Secondary Findings Section ###
if len(secondary_variants) > 0:
    doc.add_heading('Secondary Findings', level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Type'
    hdr_cells[1].text = 'Zygosity'
    hdr_cells[2].text = 'Variant/Size'
    hdr_cells[3].text = 'Gene'
    hdr_cells[4].text = 'Class'
    hdr_cells[5].text = 'Disease'

    for _, row in pd.concat([secondary_variants, empty_row], ignore_index=True).iterrows():
        r = table.add_row().cells
        chr = str(row.get("Chromosome", ""))
        start = str(row.get("Start", ""))
        end = str(row.get("End", ""))
        type = str(row.get("Type", ""))
        classification = str(row.get("Classification", ""))
        genes = str(row.get("Genes", ""))
        omim_genes = str(row.get("OMIM_gene", ""))
        zygosity = str(row.get("Zygosity", ""))
        omim = str(row.get("OMIM", ""))
        size = str(row.get("Length", ""))
        cytoband = str(row.get("Cytoband", ""))

        genes_list = genes.split(", ")
        omim_genes_list = omim_genes.split(", ")
        
        r[0].text = type
        r[1].text = zygosity
        r[2].text = f"chr{chr}:{start}-{end}" + f'\n"{cytoband}"' + f"\n({size} bp)"
        r[3].text = genes
        r[4].text = classification
        r[5].text = "- " + str(omim).replace(" | ", "\n- ") + "\n"
        for i in range(len(genes_list)):
            add_hyperlink(r[5].paragraphs[-1], f"https://www.omim.org/entry/{omim_genes_list[i]}", f"{genes_list[i]} (OMIM: {omim_genes_list[i]}) ,")
            
    last_row = table.rows[len(table.rows) - 1]
    table._tbl.remove(last_row._tr)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    center_table_content(table)

### Attached Findings Section ###
if len(attached_variants) > 0:
    doc.add_heading('Attached Findings', level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Type'
    hdr_cells[1].text = 'Zygosity'
    hdr_cells[2].text = 'Variant/Size'
    hdr_cells[3].text = 'Gene'
    hdr_cells[4].text = 'Class'
    hdr_cells[5].text = 'Disease'

    for _, row in pd.concat([attached_variants, empty_row], ignore_index=True).iterrows():
        r = table.add_row().cells
        chr = str(row.get("Chromosome", ""))
        start = str(row.get("Start", ""))
        end = str(row.get("End", ""))
        type = str(row.get("Type", ""))
        classification = str(row.get("Classification", ""))
        genes = str(row.get("Genes", ""))
        omim_genes = str(row.get("OMIM_gene", ""))
        zygosity = str(row.get("Zygosity", ""))
        omim = str(row.get("OMIM", ""))
        size = str(row.get("Length", ""))
        cytoband = str(row.get("Cytoband", ""))

        genes_list = genes.split(", ")
        omim_genes_list = omim_genes.split(", ")
        
        r[0].text = type
        r[1].text = zygosity
        r[2].text = f"chr{chr}:{start}-{end}" + f'\n"{cytoband}"' + f"\n({size} bp)"
        r[3].text = genes
        r[4].text = classification
        r[5].text = "- " + str(omim).replace(" | ", "\n- ") + "\n"
        for i in range(len(genes_list)):
            add_hyperlink(r[5].paragraphs[-1], f"https://www.omim.org/entry/{omim_genes_list[i]}", f"{genes_list[i]} (OMIM: {omim_genes_list[i]}) ,")
            
    last_row = table.rows[len(table.rows) - 1]
    table._tbl.remove(last_row._tr)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    center_table_content(table)



### Biological Interpretation of Primary Variants Section ###
if len(primary_variants) > 0:
    doc.add_heading('Biological Interpretation of Primary Findings', level=1)
    for _, row in primary_variants.iterrows():

        chr = str(row.get("Chromosome", ""))
        start = str(row.get("Start", ""))
        end = str(row.get("End", ""))
        type = str(row.get("Type", ""))
        classification = str(row.get("Classification", ""))
        genes = str(row.get("Genes", ""))
        zygosity = str(row.get("Zygosity", ""))
        omim = str(row.get("OMIM", ""))
        size = str(row.get("Length", ""))
        copy_number = str(float(str(row.get("Dosage", "")))*2)
        diseases_description = str(row.get("Diseases_description"))
        gene_description = str(row.get("Gene_description"))

        variant_paragraph = f'The {type} variant with genomic coordinates {chr}:{start}-{end} (size: ~{size} bp) was detected in {zygosity} state (copy number: ~{copy_number}). Based on the start and end coordinates (approximate) of the CNV, computed using NGS, it is located in gene(s) {genes}. Based on the 2019 ACMG guidelines, we classify this change as a "{classification}" CNV.'

        doc.add_heading(f"[CNV] {chr}:{start}-{end} in gene(s) {genes}", level=2)

        doc.add_paragraph(variant_paragraph)
        
        for disease in diseases_description.split(" | "):
            doc.add_paragraph(f"** OMIM description for " + disease)
        
        for gene in gene_description.split(" | "):
            doc.add_paragraph(f"** NCBI Gene Summary: {gene}")

### Biological Interpretation of Secondary Variants Section ###
if len(secondary_variants) > 0:
    doc.add_heading('Biological Interpretation of Secondary Findings', level=1)
    for _, row in secondary_variants.iterrows():

        chr = str(row.get("Chromosome", ""))
        start = str(row.get("Start", ""))
        end = str(row.get("End", ""))
        type = str(row.get("Type", ""))
        classification = str(row.get("Classification", ""))
        genes = str(row.get("Genes", ""))
        zygosity = str(row.get("Zygosity", ""))
        omim = str(row.get("OMIM", ""))
        size = str(row.get("Length", ""))
        copy_number = str(float(str(row.get("Dosage", "")))*2)
        diseases_description = str(row.get("Diseases_description"))
        gene_description = str(row.get("Gene_description"))

        variant_paragraph = f'The {type} variant with genomic coordinates {chr}:{start}-{end} (size: ~{size} bp) was detected in {zygosity} state (copy number: ~{copy_number}). Based on the start and end coordinates (approximate) of the CNV, computed using NGS, it is located in gene(s) {genes}. Based on the 2019 ACMG guidelines, we classify this change as a "{classification}" CNV.'

        doc.add_heading(f"[CNV] {chr}:{start}-{end} in gene(s) {genes}", level=2)

        doc.add_paragraph(variant_paragraph)
        
        for disease in diseases_description.split(" | "):
            doc.add_paragraph(f"** OMIM description for " + disease)
        
        for gene in gene_description.split(" | "):
            doc.add_paragraph(f"** NCBI Gene Summary: {gene}")

### Biological Interpretation of Attached Variants Section ###
if len(attached_variants) > 0:
    doc.add_heading('Biological Interpretation of Attached Findings', level=1)
    for _, row in attached_variants.iterrows():

        chr = str(row.get("Chromosome", ""))
        start = str(row.get("Start", ""))
        end = str(row.get("End", ""))
        type = str(row.get("Type", ""))
        classification = str(row.get("Classification", ""))
        genes = str(row.get("Genes", ""))
        zygosity = str(row.get("Zygosity", ""))
        omim = str(row.get("OMIM", ""))
        size = str(row.get("Length", ""))
        copy_number = str(float(str(row.get("Dosage", "")))*2)
        diseases_description = str(row.get("Diseases_description"))
        gene_description = str(row.get("Gene_description"))

        variant_paragraph = f'The {type} variant with genomic coordinates {chr}:{start}-{end} (size: ~{size} bp) was detected in {zygosity} state (copy number: ~{copy_number}). Based on the start and end coordinates (approximate) of the CNV, computed using NGS, it is located in gene(s) {genes}. Based on the 2019 ACMG guidelines, we classify this change as a "{classification}" CNV.'

        doc.add_heading(f"[CNV] {chr}:{start}-{end} in gene(s) {genes}", level=2)

        doc.add_paragraph(variant_paragraph)
        
        for disease in diseases_description.split(" | "):
            doc.add_paragraph(f"** OMIM description for " + disease)
        
        for gene in gene_description.split(" | "):
            doc.add_paragraph(f"** NCBI Gene Summary: {gene}")


# 4. Additional Fields
doc.add_heading('Methodology', level=1)
doc.add_paragraph("1. DNA extraction and quantification.\n2. Preparation of a genomic DNA library.\n3. Capture of the target regions by hybridization with probes in solution.\n4. Clonal amplification and sequencing of the selected gene through paired-end strategy.\n5. The average read depth of this sample is above 40x.\n6. Bioinformatics study of the DNA sequence obtained by comparison with a reference genomic sequence (GRCh38). Target regions include exons and adjacent intronic regions (±20bp) of the gene analysed.\n7. According to the information available about the patient, our diagnostic algorithm consists of:\n    - Selection of gene associated with the clinical phenotype of the patient:\n    - Study of point mutations and small InDels:\n a) Analysis of variants described in ClinVar (Pathogenic and probably pathogenic) or with a deleterious effect (frameshift, stop codon, nonsense, essential splicing, etc.).\n b) This analysis considers those variants with a number of readings >= 10x and a variant/reading ratio >0,2.\n c) Confirmation by Sanger sequencing of the pathogenic and probably pathogenic indels longer than 1bp, with low quality value and/or located in complex regions.\n    - Screening of CNVs:\n a) Clinical evaluation of CNVs identified in the gene analysed. Only those variants in genes whose associated phenotypes could be compatible with the suspected diagnosis in the patient will be reported. The CNVs analyzed will be classified into one of the following categories: pathogenic, probably pathogenic, variant of uncertain significance, probably benign or benign. This classification will be made based on the information available and taking into account the recommendations proposed in the ACMG guidelines. Probably benign and benign CNVs will not be reported.\n b) Automatic selection of control samples according to the coverage characteristics of the sample under study (10-20 control samples with an R2>= 0.9).\n c) Normalisation of the data of each sample based on the size of the fragment library and other biases.\n d) Comparison of the analysed sample with the reference samples and copy number calling.\n e) Annotation of candidate CNVs if:\n * It has a Copy Number value <= 1.5 (deletion) or >= 2.5 (duplication).\n * It includes >= 4 bins (panel) / 6 bins (exome).\n * It has a CNid score value >= 5.\n * The proportion of conflicting windows is <50%.\nOnly those CNVs with a statistical significance lower than p<0.05 (False Discovery Rate<5%) are reported. Those CNVs that do not meet these criteria are not analysed.")

doc.add_heading('Technical Limitations', level=1)
doc.add_paragraph("- In this analysis there may be regions with coverage less than the limit indicated above. The regions with special sensitivity to coverage decrease are those with high complexity (high CG content, repetitive regions, etc.).\n- The NGS technology based on the capture/enrichment through hybridization probes is not able to distinguish between high homology regions (homologous gene, pseudogenes, gene families, etc.) and may lead to false positives / negatives.\n- This study is based on the sequencing of the coding regions and adjacent intronic areas. Large deletions and rearrangements, or deep intronic mutations, would not be detected with this methodology.\n- This study is focused on the analysis of monogenic diseases in which it is possible to establish a cause-effect relationship between the presence of a variant and the development of the disease. Haplotypes or risk polymorphisms associated with diseases require other factors for their development, so they are beyond the scope of this analysis and will not be reported.\n- This methodology does not allow the detection of low proportion-mosaicisms.\n- The genes analysed in this study are directly related to the patient's clinical information provided by the physician.\n- The CNVs present in this screening report must be confirmed using an additional methodology for diagnostic purposes.\n- The NGS methodology does not allow the detection of mosaicisms, balanced chromosomal rearrangements or nucleotide expansions.\n- The chromosomal location reported refers to the CNV detected by NGS data. The technology used, based on capture probes, only covers exonic regions and essential splicing regions, so the real size of the CNV may differ from that reported in this study, as well as the break points involved in the formation of the CNV.\n- The number of copies identified in conflict regions can lead to false positives / negatives, therefore they are not reported. These results are due to variations in coverage in regions with a high proportion of conflictive windows (> 50% of windows). We consider a window as conflictive if it has overlap with a repetitive genomic element, with pseudogenes or regions of high homology, low coverage in all control samples, GC content lower than 20% or higher than 70%.")

doc.add_heading('Notes', level=1)
doc.add_paragraph("1. We assumed the existence of informed consent.\n2. Patient details included in this study are confidential and must be handled with strict confidentiality criteria.\n3. Clinical and biological interpretations are based on the bibliography currently available.\n4. The results obtained in this report should be considered in the context of other clinical and pathological factors of the patient.\n5. Paragraphs starting with ** are auto-generated and may not be available for some genes/diseases.")

doc.add_heading('Filters Used', level=1)
doc.add_paragraph("1. Read depth >= 20\n2. Variant frequency in gnomAD < 1%\n3. Gene is associated with OMIM disease.\n4. ACMG Class I, II, or III (P, LP, VUS).\n5. For CNVs: Some breakpoints affect a gene that is associated with OMIM disease")

# Save the document
doc.save(output_doc)
print(f"DOCX report has been saved to {output_doc}")
